import easyocr
import os
from docx import Document
from app.ExportacionPDF import convertir_docx_a_pdf
def conversionformato(texto_extraido,name,ruta_guardado,indicador):
    # Crear un nuevo documento
    documento = Document()
    
    # Asegurarse de que texto_extraido sea una lista y agregar cada elemento como párrafo
    if isinstance(texto_extraido, list):
        for linea in texto_extraido:
            documento.add_paragraph(linea)
    else:
        lineas = texto_extraido.split('\n')
        for linea in lineas:
            documento.add_paragraph(linea)
    

    # Guardar el documento como archivo .docx
    documento.save(ruta_guardado)
    print(f"Archivo .docx creado exitosamente en: {ruta_guardado}")

    if indicador==0:
        direccion = os.path.join(os.path.dirname(__file__), "..", "documentos")
        archivo_pdf = os.path.join(direccion, f"{os.path.splitext(name)[0]}.pdf")
        convertir_docx_a_pdf(ruta_guardado,archivo_pdf)
        
        if os.path.exists(ruta_guardado):
            os.remove(ruta_guardado)
            print(f"Archivo .docx eliminado: {ruta_guardado}")
        else:
            print("El archivo .docx no existe o ya fue eliminado.")

def extraccion(imagen,name,indicador):
    # Inicializa el lector
    lector = easyocr.Reader(['es'])  # 'es' para español
    # Lee texto de la imagen
    texto_extraido = lector.readtext(imagen, detail=0)
    # Muestra el texto extraído
    print("\n".join(texto_extraido))
    
    # Construir la ruta relativa a la carpeta "documentos"
    ruta_base = os.path.dirname(__file__)  # Directorio del archivo actual
    ruta_documentos = os.path.join(ruta_base, "..", "documentos")
    
    # Especificar el archivo .txt donde se guardará el texto
   # ruta_guardado_txt = os.path.join(ruta_documentos, "texto_extraido.txt")
    ruta_guardado = os.path.join(ruta_documentos, f"{name}.docx")
    
    # Guardar en un archivo
    try:
        # Guardar como archivo .txt
  #      with open(ruta_guardado_txt, 'w', encoding='utf-8') as archivo:
     #       archivo.write("\n".join(texto_extraido))
      #  print(f"El archivo .txt se ha guardado correctamente en: {ruta_guardado_txt}")
        
        # Guardar como archivo .docx
        conversionformato(texto_extraido,name,ruta_guardado,indicador)
    except Exception as e:
        print(f"Error al guardar el archivo: {e}")
