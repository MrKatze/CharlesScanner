import os
import tempfile
import pytest
from docx import Document
from app.ExportacionPDF import convertir_docx_a_pdf  # Ajusta el import a tu estructura real

def test_convertir_docx_a_pdf():
    # Crear archivo docx temporal
    docx_fd, docx_path = tempfile.mkstemp(suffix=".docx")
    pdf_fd, pdf_path = tempfile.mkstemp(suffix=".pdf")

    os.close(docx_fd)
    os.close(pdf_fd)

    try:
        # Crear documento DOCX de prueba
        document = Document()
        document.add_heading("Título de prueba", 0)
        document.add_paragraph("Este es un párrafo de prueba.")
        document.save(docx_path)

        # Ejecutar la conversión
        convertir_docx_a_pdf(docx_path, pdf_path)

        # Verificar que el archivo PDF fue generado y tiene contenido
        assert os.path.exists(pdf_path), "El archivo PDF no fue creado."
        assert os.path.getsize(pdf_path) > 0, "El archivo PDF está vacío."

    finally:
        # Limpiar archivos temporales
        if os.path.exists(docx_path):
            os.remove(docx_path)
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
